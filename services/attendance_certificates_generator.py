from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta
from PyPDF2 import PdfMerger
import subprocess
import os
import platform
from docx2pdf import convert



def generate_attendance_certificate(participant, formation):
    doc_path = f"static/{formation['categorie']}_ATTESTATION_DE_PARTICIPATION_A_UN_PROGRAMME_DE_DPC.docx"
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc = Document(doc_path)
    formation_titre = formation['titre'].replace('\n', '')
    formation_orientation = formation['orientation']
    formation_date_debut = formation.get('date_debut').strftime('%d/%m/%Y')
    if formation.get('date_fin'):
        formation_date_fin = formation.get('date_fin').strftime('%d/%m/%Y')
        date_signature = formation.get('date_fin') + timedelta(days=1)
    else:
        formation_date_fin = formation_date_debut
        date_signature = formation.get('date_debut') + timedelta(days=1)

    

    info_to_add = {
        "Nom :": participant['nom'].upper(),
        "Prénom :": participant['prenoms'].upper(),
        "Adresse électronique :": participant['email'],
        "N° RPPS :": participant['rpps'],
        "DD/DD/DDDD": formation_date_debut,
        "FF/FF/FFFF": formation_date_fin,
        "Année(s) civile(s) de participation :": datetime.now().year,
        "Intitulé du programme :": formation_titre,
        "Orientation nationale dans laquelle le programme s’inscrit :": formation_orientation,
        "Nom/sigle :": formation['nom'],
        "Adresse :": formation['adresse'],
        "N° enregistrement OGDPC / Agence nationale du DPC :": formation['numero dpc'],
        "//2024": date_signature.strftime("%d/%m/%Y")
    }

    for paragraph in doc.paragraphs:
        for placeholder, actual_value in info_to_add.items():
            
            if placeholder in paragraph.text:
                if placeholder == "//2024":
                    paragraph.text = paragraph.text.replace(placeholder, actual_value)
                elif placeholder == "N° RPPS :":
                    # Séparer le texte en deux parties, avant et après le placeholder
                    parts = paragraph.text.split(placeholder)
                
                    # Nettoyer le paragraphe
                    paragraph.clear()

                    # Ajouter la première partie (avant le placeholder)
                    if parts[0]:
                        paragraph.add_run(parts[0])


                    # Ajouter le placeholder sans gras
                    paragraph.add_run(placeholder)

                    # Ajouter le numéro RPPS en gras
                    actual_value = str(actual_value)
                    run = paragraph.add_run(actual_value)
                    run.bold = True

                    # Ajouter le reste du texte sans gras (après le numéro RPPS)
                    remaining_text = parts[1].replace(actual_value, '')
                    paragraph.add_run(remaining_text)
                elif placeholder == "DD/DD/DDDD" or placeholder == "FF/FF/FFFF":
                    # Si le paragraphe contient des dates à remplacer

                    # Remplacer d'abord la date de début, puis la date de fin en nettoyant le texte mais en réécrivant proprement
                    text_parts = paragraph.text.split("DD/DD/DDDD")

                    # Nettoyer le paragraphe
                    paragraph.clear()

                    # Ajouter le texte avant la date de début, puis la date de début en gras
                    paragraph.add_run(text_parts[0])
                    run = paragraph.add_run(formation_date_debut)
                    run.bold = True

                    # Si la phrase continue avec "Date de fin :", traiter cela
                    if len(text_parts) > 1:
                        remaining_text = text_parts[1].split("FF/FF/FFFF")
                        
                        # Ajouter le texte entre les deux dates (souvent " Date de fin :")
                        paragraph.add_run(remaining_text[0])

                        # Ajouter la date de fin en gras
                        run = paragraph.add_run(formation_date_fin)
                        run.bold = True

                        # Ajouter tout texte restant après la date de fin
                        if len(remaining_text) > 1:
                            paragraph.add_run(remaining_text[1])
                else:
                    parts = paragraph.text.split(placeholder)
                    paragraph.text = parts[0] + placeholder
                    paragraph.add_run(f" {actual_value}").bold = True
                    if len(parts) > 1:
                        paragraph.add_run(parts[1]).bold = True
                # Vérifier si le paragraphe contient "Cachet et signature"
        if "Cachet et signature" in paragraph.text:
            # Ajouter l'image juste après ce paragraphe
            image_path = os.path.join(script_dir, "../static/ekoforma_stamp.png")
            # Effacer le texte existant
            paragraph.clear()
            # Réécrire "Cachet et signature"
            paragraph.add_run("Cachet et signature")
            # Ajouter l'image immédiatement après le texte
            paragraph.add_run("\n")
            paragraph.add_run().add_picture(image_path, width=Inches(1.5))  # Ajustez la taille de l'image selon vos besoins
            # Centrer le texte et l'image
            paragraph.alignment = 2
    output_directory = os.path.join(script_dir, "../downloads")
    output_file = os.path.join(output_directory, f"../downloads/{formation.get('code')}_ATTESTATION_DE_PARTICIPATION_A_UN_PROGRAMME_DE_DPC_{participant['nom_complet']}.docx")                        

    doc.save(output_file)

    # Conversion du fichier Word en PDF en utilisant LibreOffice en mode headless
    pdf_output_file = output_file.replace(".docx", ".pdf")
    

    system_platform = platform.system()

    if system_platform == 'Linux' or system_platform == 'Darwin':
        try:
            # Commande pour convertir .docx en .pdf
            command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_directory, output_file]
            subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            print(f"PDF généré avec succès : {pdf_output_file}")
        except subprocess.CalledProcessError as e:
            print(f"Erreur lors de la conversion en PDF : {e.stderr.decode('utf-8')}")
    elif system_platform == 'Windows':
        try:
            convert(output_file, output_directory)
            print(f"PDF généré avec succès : {pdf_output_file}")
        except Exception as e:
            print(f"Erreur lors de la conversion en PDF (docx2pdf) : {str(e)}")
    else:
        print(f"Système d'exploitation {system_platform} non pris en charge pour cette conversion.")
    
    return pdf_output_file


def merge_pdfs(formation_code, output_directory):
    merger = PdfMerger()
    for filename in os.listdir(output_directory):
        if filename.endswith(".pdf") and formation_code in filename and 'merged' not in filename:
            file_path = os.path.join(output_directory, filename)
            print(filename)
            print('----')
            merger.append(file_path)
    
    merged_pdf_path = os.path.join(output_directory, f"{formation_code}_merged_document.pdf")
    with open(merged_pdf_path, 'wb') as merged_file:
        merger.write(merged_file)

    print(f"PDF fusionné généré : {merged_pdf_path}")

    return merged_pdf_path
